from docx import Document
from tkinter import Tk, filedialog, messagebox, Toplevel, Text, Scrollbar, Button
from tkinter import ttk
import os
from pathlib import Path
import logging

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(message)s')

def get_numbers_from_file(file_path):
    path = Path(file_path)
    if not path.exists():
        messagebox.showerror("Error", f"Файл {file_path} не найден. Для работы программы создайте {file_path} рядом с файлом программы")
        exit(1)

    with path.open('r') as file:
        numbers = [line.strip().replace(',', '.') for line in file.readlines()]
    return numbers

def copy_formatting(src_run, dest_run):
    dest_run.bold = src_run.bold
    dest_run.italic = src_run.italic
    dest_run.underline = src_run.underline
    dest_run.font.size = src_run.font.size
    dest_run.font.name = src_run.font.name
    if src_run.font.color.rgb:
        dest_run.font.color.rgb = src_run.font.color.rgb

def update_table_cells(table, numbers_set, found_numbers):
    for row in table.rows:
        cell_text = row.cells[1].text.strip().replace(',', '.')
        if cell_text in numbers_set:
            old_text_run = row.cells[4].paragraphs[0].runs[0]
            row.cells[4].text = ''
            new_run = row.cells[4].paragraphs[0].add_run('Транспорт')
            copy_formatting(old_text_run, new_run)
            found_numbers[cell_text] = True

def update_word_file(doc_path, numbers):
    doc = Document(doc_path)
    numbers_set = set(numbers)
    found_numbers = {num: False for num in numbers}

    for table in doc.tables:
        update_table_cells(table, numbers_set, found_numbers)

    doc_dir = Path(doc_path).parent
    new_doc_path = doc_dir / (Path(doc_path).stem + " с транспортом.docx")
    doc.save(new_doc_path)
    logging.info(f"Документ сохранен как {new_doc_path}")
    return found_numbers

def log_replacement_results(found_numbers):
    results = []
    for number, found in found_numbers.items():
        if found:
            msg = f"{number}: Заменено"
            logging.info(msg)
            results.append(msg)
        else:
            msg = f"{number}: Не найдено"
            logging.info(msg)
            results.append(msg)
    return results

def show_log_window(results):
    log_window = Toplevel()
    log_window.title("Результаты замены")
    log_window.geometry("400x300")

    text_area = Text(log_window, wrap='word')
    text_area.pack(expand=True, fill='both')

    scrollbar = Scrollbar(text_area)
    scrollbar['command'] = text_area.yview
    text_area['yscrollcommand'] = scrollbar.set
    scrollbar.pack(side='right', fill='y')

    for result in results:
        text_area.insert('end', result + '\n')

    text_area.config(state='disabled')

    close_button = Button(log_window, text="Закрыть", command=log_window.destroy)
    close_button.pack(pady=5)

    log_window.mainloop()  # Используем mainloop для поддержания окна открытым

def main():
    root = Tk()
    root.withdraw()

    numbers = get_numbers_from_file('numbers.txt')

    doc_path = filedialog.askopenfilename(title="Выберите документ Word", filetypes=[("Word files", "*.docx")])

    if doc_path:
        found_numbers = update_word_file(doc_path, numbers)
        results = log_replacement_results(found_numbers)
        show_log_window(results)
    else:
        logging.warning("Файл не выбран.")

if __name__ == "__main__":
    # Вывод id по которому можно найти программу
    print("id 1_1")
    main()
